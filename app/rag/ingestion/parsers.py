"""
Document parser for Enterprise AI Platform.

Google Cloud Storage is the single source of truth.
"""

import json
from io import BytesIO
from pathlib import Path

from docx import Document
from google.cloud import storage
from openpyxl import load_workbook
from pypdf import PdfReader


class DocumentParser:
    """
    Parses enterprise documents stored in Google Cloud Storage.
    """

    def __init__(self) -> None:

        self._storage_client = storage.Client()

    def parse(
        self,
        source_path: str,
    ) -> str:
        """
        Parse a document from Google Cloud Storage.

        Args:
            gcs_uri:
                gs://bucket/file.pdf

        Returns:
            Extracted text.
        """

        file_bytes, extension = self._download_from_gcs(
            source_path
        )

        if extension == ".pdf":
            return self._extract_pdf(file_bytes)

        elif extension == ".txt":
            return self._extract_txt(file_bytes)

        elif extension == ".docx":
            return self._extract_docx(file_bytes)

        elif extension == ".xlsx":
            return self._extract_xlsx(file_bytes)

        elif extension == ".json":
            return self._extract_json(file_bytes)

        raise ValueError(
            f"Unsupported file type: {extension}"
        )

    def _download_from_gcs(
        self,
        source_path: str,
    ) -> tuple[bytes, str]:
        """
        Download a document from Google Cloud Storage.
        """

        path = source_path.replace(
            "gs://",
            "",
        )

        bucket_name, blob_name = path.split(
            "/",
            1,
        )

        bucket = self._storage_client.bucket(
            bucket_name
        )

        blob = bucket.blob(
            blob_name
        )

        file_bytes = blob.download_as_bytes()

        extension = Path(
            blob_name
        ).suffix.lower()

        return file_bytes, extension

    def _extract_pdf(
        self,
        file_bytes: bytes,
    ) -> str:
        """
        Extract text from PDF.
        """

        reader = PdfReader(
            BytesIO(file_bytes)
        )

        text = ""

        for page in reader.pages:

            page_text = page.extract_text()

            if page_text:
                text += page_text + "\n"

        return text

    def _extract_txt(
        self,
        file_bytes: bytes,
    ) -> str:
        """
        Extract text from TXT.
        """

        return file_bytes.decode(
            "utf-8",
            errors="ignore",
        )
    
    def _extract_docx(
        self,
        file_bytes: bytes,
    ) -> str:
        """
        Extract text from DOCX.
        """

        document = Document(
            BytesIO(file_bytes)
        )

        texts = []

        #
        # Paragraphs
        #

        for paragraph in document.paragraphs:

            if paragraph.text.strip():

                texts.append(
                    paragraph.text
                )

        #
        # Tables
        #

        for table in document.tables:

            for row in table.rows:

                row_text = " | ".join(
                    cell.text.strip()
                    for cell in row.cells
                )

                texts.append(
                    row_text
                )

        return "\n".join(texts)

    def _extract_xlsx(
        self,
        file_bytes: bytes,
    ) -> str:
        """
        Extract text from Excel.
        """

        workbook = load_workbook(
            BytesIO(file_bytes),
            data_only=True,
        )

        text_parts = []

        for sheet in workbook.worksheets:

            text_parts.append(
                f"\n=== Sheet: {sheet.title} ===\n"
            )

            rows = list(
                sheet.iter_rows(
                    values_only=True
                )
            )

            if not rows:
                continue

            headers = [
                str(header).strip()
                if header is not None
                else f"Column_{index}"
                for index, header in enumerate(
                    rows[0],
                    start=1,
                )
            ]

            for row in rows[1:]:

                row_lines = []

                for header, value in zip(
                    headers,
                    row,
                ):

                    if value is not None:

                        row_lines.append(
                            f"{header}: {value}"
                        )

                if row_lines:

                    text_parts.append(
                        "\n".join(row_lines)
                    )

                    text_parts.append(
                        "\n--------------------\n"
                    )

        return "\n".join(text_parts)

    def _extract_json(
        self,
        file_bytes: bytes,
    ) -> str:
        """
        Extract text from JSON.
        """

        data = json.loads(
            file_bytes.decode(
                "utf-8"
            )
        )

        return "\n".join(
            self._flatten_json(data)
        )

    def _flatten_json(
        self,
        obj,
        prefix: str = "",
    ) -> list[str]:
        """
        Flatten nested JSON.
        """

        lines = []

        if isinstance(obj, dict):

            for key, value in obj.items():

                new_prefix = (
                    f"{prefix}.{key}"
                    if prefix
                    else key
                )

                lines.extend(
                    self._flatten_json(
                        value,
                        new_prefix,
                    )
                )

        elif isinstance(obj, list):

            for index, item in enumerate(obj):

                lines.extend(
                    self._flatten_json(
                        item,
                        f"{prefix}[{index}]",
                    )
                )

        else:

            lines.append(
                f"{prefix}: {obj}"
            )

        return lines 